import fitz  # PyMuPDF
import io
import json
import base64
import os
from PIL import Image
from openai import OpenAI
from google import genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor

def extract_images_from_pdf(pdf_bytes, start_page=1, end_page=None):
    """从上传的 PDF 字节流中提取指定范围内的所有图片（含矢量图页面渲染）"""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    extracted_images = []
    
    total_pages = len(doc)
    if end_page is None or end_page > total_pages:
        end_page = total_pages
        
    for page_idx in range(start_page - 1, end_page):
        page = doc[page_idx]
        # 1. 尝试提取内嵌的位图图片
        image_list = page.get_images(full=True)
        page_has_raster = False
        
        for img_idx, img_meta in enumerate(image_list):
            xref = img_meta[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            
            image = Image.open(io.BytesIO(image_bytes))
            if image.width > 150 and image.height > 150:
                # CMYK/Palette modes must be converted to RGB before saving to PNG
                if image.mode in ("CMYK", "P"):
                    img_to_save = image.convert("RGB")
                else:
                    img_to_save = image
                
                png_stream = io.BytesIO()
                img_to_save.save(png_stream, format="PNG")
                png_bytes = png_stream.getvalue()
                
                extracted_images.append({
                    "page": page_idx + 1,
                    "index": img_idx + 1,
                    "bytes": png_bytes,
                    "format": "png",
                    "image_obj": image,
                    "is_full_page": False
                })
                page_has_raster = True
                
        # 2. 兜底机制：如果页面中没有内嵌位图，但页面中含有图表关键字
        if not page_has_raster:
            text = page.get_text().lower()
            if any(kw in text for kw in ["figure", "fig.", "图", "scatterplot", "bar chart", "histogram", "boxplot"]):
                # 渲染页面为高分辨率图片 (2.0 倍缩放，确保大模型能看清字)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                image_bytes = pix.tobytes("png")
                image = Image.open(io.BytesIO(image_bytes))
                extracted_images.append({
                    "page": page_idx + 1,
                    "index": 1,
                    "bytes": image_bytes,
                    "format": "png",
                    "image_obj": image,
                    "is_full_page": True
                })
                
    return extracted_images

import hashlib

CACHE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".cache")

def get_image_dhash(image_bytes):
    """计算图像的知觉哈希（Difference Hash, dHash）"""
    try:
        img = Image.open(io.BytesIO(image_bytes))
        # 缩放到 9x8 并转换为灰度图
        img_gray = img.convert('L').resize((9, 8), Image.Resampling.LANCZOS)
        pixels = list(img_gray.getdata())
        
        diff = []
        for row in range(8):
            for col in range(8):
                pixel_left = pixels[row * 9 + col]
                pixel_right = pixels[row * 9 + col + 1]
                diff.append(pixel_left > pixel_right)
                
        # 转换为 16 位十六进制字符串
        decimal_val = 0
        hash_str = ""
        for index, val in enumerate(diff):
            if val:
                decimal_val += 2**(index % 8)
            if index % 8 == 7:
                hash_str += f"{decimal_val:02x}"
                decimal_val = 0
        return hash_str
    except Exception:
        return ""

def hamming_distance(hash1, hash2):
    """计算两个十六进制 dHash 之间的汉明距离（不同位数的个数）"""
    try:
        h1 = int(hash1, 16)
        h2 = int(hash2, 16)
        return bin(h1 ^ h2).count('1')
    except Exception:
        return 999

def get_cached_result(image_bytes, language):
    """使用 MD5（精确匹配）与 dHash 汉明距离（结构相似匹配）查找缓存"""
    if not os.path.exists(CACHE_DIR):
        os.makedirs(CACHE_DIR)
        
    img_hash = hashlib.md5(image_bytes).hexdigest()
    
    # 1. 精确匹配：文件存在则直接返回
    exact_cache_path = os.path.join(CACHE_DIR, f"{img_hash}_{language}.json")
    if os.path.exists(exact_cache_path):
        try:
            with open(exact_cache_path, "r", encoding="utf-8") as f:
                return json.load(f)["data"]
        except Exception:
            pass
            
    # 2. 相似度匹配：扫描缓存目录进行 dHash 对比
    new_dhash = get_image_dhash(image_bytes)
    if not new_dhash:
        return None
        
    for filename in os.listdir(CACHE_DIR):
        if filename.endswith(f"_{language}.json"):
            cache_path = os.path.join(CACHE_DIR, filename)
            try:
                with open(cache_path, "r", encoding="utf-8") as f:
                    entry = json.load(f)
                cached_dhash = entry.get("dhash", "")
                if cached_dhash:
                    dist = hamming_distance(new_dhash, cached_dhash)
                    # 汉明距离 <= 6 认为图表结构高度相似 (64位哈希中只有不到10%不同)
                    if dist <= 6:
                        # 复用结果，并为当前新图片写入精确缓存以提高下次加载效率
                        save_cache_result(image_bytes, language, entry["data"])
                        return entry["data"]
            except Exception:
                continue
    return None

def save_cache_result(image_bytes, language, data):
    """保存结果到缓存，写入 MD5 与 dHash"""
    if not os.path.exists(CACHE_DIR):
        os.makedirs(CACHE_DIR)
        
    img_hash = hashlib.md5(image_bytes).hexdigest()
    dhash = get_image_dhash(image_bytes)
    
    cache_path = os.path.join(CACHE_DIR, f"{img_hash}_{language}.json")
    cache_entry = {
        "md5": img_hash,
        "dhash": dhash,
        "data": data
    }
    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump(cache_entry, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def analyze_chart_accessibility(image_bytes, is_full_page=False, language="简体中文"):
    """调用大模型（OpenAI GPT-4o 或 Google Gemini 2.5 Flash）提取学术图表的无障碍要素与数据表格"""
    # 1. 检查本地缓存以提升分析速度（支持精准匹配与结构感知匹配）
    cached_data = get_cached_result(image_bytes, language)
    if cached_data:
        return cached_data

    api_key = os.environ.get("OPENAI_API_KEY")
    
    system_prompt = f"""You are a digital accessibility and information architecture expert. 
Please analyze the academic chart and output a JSON object in the following format (no markdown code blocks):
{{
  "alt_text": "Single-sentence description explaining the chart type and core subject (under 50 words), written in {language}.",
  "trend_summary": "Analysis of academic trends, extreme values, and key inflection points (under 150 words), written in {language}.",
  "table_headers": ["Column 1", "Column 2", ... (Headers translated or written in {language})],
  "table_rows": [
    ["Data 1", "Data 2", ...],
    ["Data 3", "Data 4", ...]
  ]
}}"""

    user_text = f"Please analyze this academic chart, extract all data, and reconstruct it into an accessible semantic structure. Output all textual descriptions and table headers in {language}."
    if is_full_page:
        user_text = f"This image is a full academic paper page containing text and figures. Please ignore the body text, locate the chart (such as a scatterplot, histogram, line graph, or boxplot), extract all of its data, and reconstruct it into an accessible semantic structure. Output all textual descriptions and table headers in {language}."

    # 判断 API Key 的类型
    is_gemini_key = api_key and (api_key.startswith("AQ.") or api_key.startswith("AIza"))
    
    if is_gemini_key:
        client = genai.Client(api_key=api_key)
        image = Image.open(io.BytesIO(image_bytes))
        
        import re
        import time
        
        def _call_gemini_model_with_dynamic_retry(model_name):
            max_attempts = 4
            for attempt in range(max_attempts):
                try:
                    return client.models.generate_content(
                        model=model_name,
                        contents=[user_text, image],
                        config=dict(
                            system_instruction=system_prompt,
                            response_mime_type="application/json"
                        )
                    )
                except Exception as e:
                    err_str = str(e)
                    # Check if it's a 429 rate limit or RESOURCE_EXHAUSTED error
                    if "429" in err_str or "RESOURCE_EXHAUSTED" in err_str:
                        # Extract the exact retry delay requested by Gemini
                        match = re.search(r"retry in (\d+\.?\d*)s", err_str, re.IGNORECASE)
                        if match:
                            delay = float(match.group(1))
                            # Add a small buffer of 1.5 seconds to be safe
                            time.sleep(delay + 1.5)
                            continue
                        else:
                            time.sleep(8 * (attempt + 1))
                            continue
                    # Check if it's a 503 unavailable error
                    elif "503" in err_str or "UNAVAILABLE" in err_str:
                        time.sleep(5 * (attempt + 1))
                        continue
                    else:
                        raise e
            # Last resort attempt
            return client.models.generate_content(
                model=model_name,
                contents=[user_text, image],
                config=dict(
                    system_instruction=system_prompt,
                    response_mime_type="application/json"
                )
            )
            
        try:
            # 1st attempt: Try stable gemini-flash-latest
            response = _call_gemini_model_with_dynamic_retry('gemini-flash-latest')
        except Exception as e:
            try:
                # 2nd attempt (fallback): Try lighter gemini-flash-lite-latest under high load
                response = _call_gemini_model_with_dynamic_retry('gemini-flash-lite-latest')
            except Exception:
                raise e
                
        result_data = json.loads(response.text)
        save_cache_result(image_bytes, language, result_data)
        return result_data
    else:
        client = OpenAI(api_key=api_key)
        base64_img = base64.b64encode(image_bytes).decode('utf-8')
        
        from tenacity import retry, stop_after_attempt, wait_exponential
        
        @retry(
            stop=stop_after_attempt(3),
            wait=wait_exponential(multiplier=1, min=2, max=10),
            reraise=True
        )
        def _call_openai_with_retry():
            return client.chat.completions.create(
                model="gpt-4o",
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": [
                        {"type": "text", "text": user_text},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_img}"}}
                    ]}
                ]
            )
            
        response = _call_openai_with_retry()
        result_data = json.loads(response.choices[0].message.content)
        save_cache_result(image_bytes, language, result_data)
        return result_data

def create_accessible_docx(processed_charts):
    """构建符合无障碍规范的 Word 文档"""
    doc = Document()
    
    # 标题
    title = doc.add_heading("学术文献图表无障碍转译报告", level=1)
    
    for idx, item in enumerate(processed_charts):
        doc.add_heading(f"图表 {idx + 1}（源自 PDF 第 {item['page']} 页）", level=2)
        
        # 1. 插入图片
        img_stream = io.BytesIO(item['bytes'])
        doc.add_picture(img_stream, width=Inches(4.5))
        
        # 2. 替代文本标签
        p_alt = doc.add_paragraph()
        p_alt_run = p_alt.add_run(f"【Alt Text / 读屏替代文本】: {item['data']['alt_text']}")
        p_alt_run.font.size = Pt(9.5)
        p_alt_run.font.italic = True
        p_alt_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # 3. 趋势概括
        doc.add_heading("核心数据趋势", level=3)
        doc.add_paragraph(item['data']['trend_summary'])
        
        # 4. 语义化数据表格（供 NVDA/JAWS 遍历）
        doc.add_heading("图表原始数据还原表", level=3)
        headers = item['data']['table_headers']
        rows = item['data']['table_rows']
        
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Shading Accent 1'
        
        # 填充表头
        for col_idx, header_text in enumerate(headers):
            cell = table.cell(0, col_idx)
            cell.text = str(header_text)
            
        # 填充数据行
        for row_idx, row_data in enumerate(rows):
            for col_idx, val in enumerate(row_data):
                table.cell(row_idx + 1, col_idx).text = str(val)
                
        doc.add_paragraph() # 空行分隔
        
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream
